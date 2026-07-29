from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from pathlib import Path

ROOT = Path(__file__).resolve().parent
OUT = ROOT / 'CodePulse_Project_Report_Chapters_1_to_3.docx'

BLUE = '1F4D78'; LIGHT = 'EAF1F7'; DARK = '163A5A'; GREY = '5A6772'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margin(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); mar = tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar = OxmlElement('w:tcMar'); tcPr.append(mar)
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        n = mar.find(qn(f'w:{side}'))
        if n is None: n = OxmlElement(f'w:{side}'); mar.append(n)
        n.set(qn('w:w'), str(val)); n.set(qn('w:type'), 'dxa')

def border_table(table, color='B9C6D2'):
    tblPr = table._tbl.tblPr; borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),color); borders.append(e)
    tblPr.append(borders)

def set_width(cell, width):
    cell.width = Inches(width)
    tcPr = cell._tc.get_or_add_tcPr(); w = tcPr.find(qn('w:tcW'))
    if w is None: w = OxmlElement('w:tcW'); tcPr.append(w)
    w.set(qn('w:w'), str(int(width*1440))); w.set(qn('w:type'),'dxa')

def add_text(cell, text, bold=False, color=None, size=9.3):
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.08
    r = p.add_run(str(text)); r.bold=bold; r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); r._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margin(cell)

def table(doc, headers, rows, widths=None, font=8.8):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False; border_table(t)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; add_text(c,h,True,'FFFFFF',font); shade(c,BLUE)
        if widths: set_width(c,widths[i])
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            add_text(cells[i],v,False,None,font)
            if widths: set_width(cells[i],widths[i])
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def para(doc, text='', style=None, boldlead=None):
    p=doc.add_paragraph(style=style)
    p.paragraph_format.space_after=Pt(7); p.paragraph_format.line_spacing=1.18
    if boldlead and text.startswith(boldlead):
        r=p.add_run(boldlead); r.bold=True; p.add_run(text[len(boldlead):])
    else: p.add_run(text)
    return p

def bullet(doc, text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.12; p.add_run(text); return p

def heading(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); p.paragraph_format.keep_with_next=True; p.add_run(text); return p

def caption(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(10)
    r=p.add_run(text); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(GREY)

def diagram(doc, title, rows):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; border_table(t, '91A9BC')
    c=t.cell(0,0); shade(c,'F6F9FC'); set_width(c,6.3); set_cell_margin(c,160,180,160,180)
    c.text=''
    for i,line in enumerate(rows):
        p=c.paragraphs[0] if i==0 else c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
        r=p.add_run(line); r.font.name='Consolas'; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(DARK); r.bold = ('[' in line or '│' not in line and '↓' not in line and '→' not in line)
    caption(doc,title)

def page_break(doc): doc.add_page_break()

doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.85); sec.right_margin=Inches(0.85)
styles=doc.styles
styles['Normal'].font.name='Calibri'; styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); styles['Normal']._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); styles['Normal'].font.size=Pt(10.5)
for s,size,color in [('Heading 1',16,BLUE),('Heading 2',13,BLUE),('Heading 3',11,DARK)]:
    styles[s].font.name='Calibri'; styles[s].font.size=Pt(size); styles[s].font.bold=True; styles[s].font.color.rgb=RGBColor.from_string(color)
    styles[s].paragraph_format.space_before=Pt(12); styles[s].paragraph_format.space_after=Pt(6)

# footer
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=footer.add_run('CodePulse Project Report | Chapters 1-3'); rr.font.size=Pt(8); rr.font.color.rgb=RGBColor.from_string(GREY)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(110)
r=p.add_run('CODEPULSE'); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string(DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Engineering Intelligence Platform'); r.font.size=Pt(17); r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(22); r=p.add_run('Project Report: Chapters 1 to 3'); r.font.size=Pt(14); r.bold=True
para(doc,'A structured report covering project introduction, system analysis and requirements, and system design.',None)
doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80); r=p.add_run('Prepared for academic submission\nJuly 2026'); r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(GREY)
page_break(doc)

# CHAPTER 1
heading(doc,'CHAPTER 1: INTRODUCTION',1)
para(doc,'This chapter introduces CodePulse, its motivation, objectives, scope, societal relevance, and report structure.')
heading(doc,'1.1 Background of the Project',2)
para(doc,'Modern software repositories accumulate code, documentation, dependency, and change-history data faster than engineering teams can continuously interpret it. Traditional static-analysis tools identify isolated code smells, but often do not connect those findings to documentation gaps, ownership concentration, change churn, or practical remediation. This makes it difficult for teams to understand where maintenance risk is concentrated and which improvements should be prioritized.')
para(doc,'CodePulse is an AI-powered Engineering Intelligence Platform designed to analyze a GitHub repository as a living engineering system. It collects repository structure, source-file metadata, documentation, commit history, and dependency relationships; uses these signals to support detection of documentation drift, technical debt, knowledge debt, and maintainability risk; and presents evidence through a dashboard and report-oriented views. Its AI Explainability Engine is designed to translate analytical evidence into human-readable explanations and actionable recommendations.')
para(doc,'The project aligns primarily with United Nations Sustainable Development Goal 9: Industry, Innovation and Infrastructure. By helping teams build more reliable, maintainable, and knowledge-resilient software, CodePulse supports digital innovation and resilient technical infrastructure. It also contributes indirectly to SDG 4 (Quality Education) by improving onboarding knowledge and SDG 12 (Responsible Consumption and Production) by reducing avoidable redevelopment and maintenance waste.')
heading(doc,'1.2 Objectives',2)
bullet(doc,'Develop a secure web platform through which authenticated users can submit and manage repository analyses.')
bullet(doc,'Extract repository metadata, files, documentation records, commit history, dependency edges, contributors, and dependency manifests from supported public GitHub repositories.')
bullet(doc,'Provide a unified dashboard that presents repository intelligence, health-oriented signals, risk-oriented views, and report-ready evidence.')
bullet(doc,'Establish an extensible analytical architecture for documentation-drift detection, technical-debt analysis, knowledge-debt assessment, risk scoring, and AI-assisted remediation.')
bullet(doc,'Persist analysis results in a structured MongoDB data model so that repository evidence can be retrieved, compared, and reported.')
heading(doc,'1.3 Purpose, Scope and Applicability',2)
heading(doc,'Purpose',3)
para(doc,'The purpose of CodePulse is to reduce the gap between raw repository data and engineering decisions. Teams need more than a list of warnings: they need a coherent view of what has changed, what is difficult to maintain, where knowledge is fragile, and what action is justified by evidence. The project improves the current fragmented workflow by combining repository inspection, documentation awareness, commit activity, dependency information, risk prioritization, and explainable recommendations in one platform.')
heading(doc,'Scope',3)
para(doc,'The current implementation focuses on the Repository Intelligence foundation: validated public GitHub URL analysis, shallow cloning, metadata extraction, documentation extraction, commit-history collection, basic JavaScript/TypeScript/Python dependency-edge extraction, MongoDB persistence, authenticated repository retrieval, and dashboard/report presentation. The intended analytical scope additionally includes knowledge drift, technical debt, knowledge debt, risk intelligence, and AI explanation services.')
para(doc,'Assumptions include access to a public GitHub repository, Git availability on the server, a configured MongoDB instance, and an internet connection for GitHub metadata and manifest retrieval. Interactive analysis is deliberately bounded by configurable clone timeout, repository-size, file-count, source-file-count, and source-file-size limits. The present public analysis endpoint accepts GitHub URLs; GitLab is supported as a connected OAuth source but is outside the current public analysis input path. Planned scoring and recommendation endpoints remain future analytical milestones rather than completed claims.')
heading(doc,'Applicability',3)
bullet(doc,'Engineering managers can prioritize maintenance investment using evidence about repository composition, activity, documentation, and dependencies.')
bullet(doc,'Developers can use repository and dependency views to understand unfamiliar codebases and identify candidates for documentation or refactoring work.')
bullet(doc,'Project teams can reduce key-person dependency by identifying concentrated contribution patterns and documentation gaps.')
bullet(doc,'Educational institutions and open-source communities can use the platform to teach software maintainability and assess repository health transparently.')
heading(doc,'1.4 Overview of the Report',2)
table(doc,['Chapter','Coverage'],[
('Chapter 1: Introduction','Project background, SDG alignment, objectives, purpose, scope, applicability, and report structure.'),
('Chapter 2: System Analysis and Requirements','Existing-system analysis, proposed system, requirements specification, user characteristics, constraints, and block/context model.'),
('Chapter 3: System Design','Architecture, modules, data flows, ER model, database design, interface/procedural design, and report design.')], [1.85,4.65],9.2)

page_break(doc)
# CHAPTER 2
heading(doc,'CHAPTER 2: SYSTEM ANALYSIS AND REQUIREMENTS',1)
para(doc,'This chapter analyzes current workflows and specifies CodePulse requirements, users, constraints, and conceptual model.')
heading(doc,'2.1 Existing System',2)
para(doc,'Engineering teams commonly use a mixture of source-control platforms, static-analysis tools, issue trackers, documentation sites, dependency viewers, and manual code reviews. GitHub provides repository browsing and history; linters and quality tools provide code-level warnings; documentation systems provide separate written knowledge. These tools are useful, but their outputs are often separated and require an engineer to manually correlate code complexity, churn, dependencies, missing documentation, and ownership patterns.')
heading(doc,'2.2 Limitations of the Existing System',2)
bullet(doc,'Repository health evidence is fragmented across multiple products and files, making prioritization slow and subjective.')
bullet(doc,'Code quality indicators may not be connected to documentation freshness or historical change activity.')
bullet(doc,'New contributors face high onboarding effort because architecture, dependencies, and documentation coverage are not consolidated.')
bullet(doc,'Generic warnings rarely explain why an issue matters in the context of a specific repository or what sequence of remediation is most useful.')
bullet(doc,'Manual collection of repository metrics is difficult to repeat consistently across projects and over time.')
heading(doc,'2.3 Proposed System',2)
para(doc,'CodePulse addresses the problem of turning scattered repository signals into prioritized, explainable engineering intelligence. An authenticated user submits a supported public GitHub repository URL. The system validates the request, inspects repository metadata, performs a bounded shallow clone, extracts files, documentation, commits and dependencies, and stores structured evidence. The dashboard then enables retrieval and review of this evidence. Planned analytical services derive debt, drift, risk, and AI-backed recommendations from the stored repository model.')
table(doc,['Sub-problem','CodePulse response'],[
('Repository visibility','Build a structured repository model from files, commits, documentation, dependencies, contributors, and manifests.'),
('Knowledge drift','Compare implementation structure and documentation to identify missing, outdated, incorrect, or dead documentation (planned analytical service).'),
('Maintainability risk','Compute complexity, duplication, circular-dependency, churn, and ownership signals, then prioritize modules (planned analytical service).'),
('Decision support','Present evidence in a responsive dashboard and produce explainable, action-oriented recommendations (AI stage planned).'),
('Secure access','Use authenticated user accounts, protected routes, secure session handling, rate limiting, and owned-resource checks.')], [1.75,4.75],8.8)
heading(doc,'2.4 Benefits of the Proposed System',2)
bullet(doc,'Consolidates repository evidence into a single engineering-intelligence workspace.')
bullet(doc,'Reduces manual effort in collecting repository, documentation, commit, and dependency information.')
bullet(doc,'Supports more informed maintenance prioritization through evidence-based views and planned risk scoring.')
bullet(doc,'Improves maintainability and onboarding by making documentation coverage and repository structure more visible.')
bullet(doc,'Provides an extensible foundation for AI-generated explanations that remain grounded in repository evidence.')
heading(doc,'2.5 Features of the Proposed System',2)
table(doc,['Feature','Description'],[
('Authentication and accounts','Email/password authentication, verification and reset flows, refresh sessions, profile/settings, GitHub and GitLab OAuth connections.'),
('Repository analysis','GitHub URL validation, metadata inspection, shallow clone, bounded file parsing, documentation extraction, commit collection and dependency-edge extraction.'),
('Evidence retrieval','Owned-repository list/detail, paginated files, commits, dependencies, documentation, contributors and manifest dependencies.'),
('Dashboard and reporting','Repository overview, pipeline status, dependency and repository-intelligence panels, technical/knowledge/risk-oriented screens and a print-ready report view.'),
('Future analytics','Knowledge drift, technical debt, knowledge debt, risk intelligence, and evidence-backed AI recommendations behind backend API boundaries.')], [1.75,4.75],8.8)
heading(doc,'2.6 System Requirements Specification',2)
para(doc,'The following requirements specify what CodePulse must provide, independent of implementation choices. Requirements marked “planned” describe the target analytical platform and are supported by the documented API contract; they are not represented as completed runtime behavior in the current foundation.')
heading(doc,'2.6.1 User Characteristics',3)
table(doc,['User type','Characteristics and permitted tasks'],[
('Engineering manager / team lead','Moderate technical knowledge; reviews repository health, prioritizes risks, reads reports and recommendations.'),
('Developer / maintainer','Technical user; submits repositories, explores files, documentation, commits, dependencies, and remediation evidence.'),
('Project administrator / account user','Manages profile, settings, authentication, connected repository sources, and owned repository records.'),
('System administrator','Maintains hosting, environment variables, MongoDB connectivity, backups, security configuration, and service availability.')], [1.8,4.7],8.8)
heading(doc,'2.6.2 Software and Hardware Requirements',3)
table(doc,['Category','Minimum / recommended requirement'],[
('Operating system','Windows 10/11, Linux, or macOS for development; Linux container/VM recommended for production.'),
('Runtime and compiler/build tools','Node.js 24-compatible runtime for production image; npm; Vite build tool; Git executable available on PATH for repository cloning.'),
('Backend / database','Express 5, MongoDB 7 driver, MongoDB server or Atlas-compatible instance, dotenv, bcryptjs, undici.'),
('Frontend libraries','React 19, React DOM, Vite 8, Tailwind CSS 4, Recharts, Lucide React, Radix Slot, class-variance-authority.'),
('Testing and quality tools','Node built-in test runner for backend fixtures; Oxlint; browser developer tools. Optional container tooling: Docker.'),
('Processor','Dual-core 2.0 GHz minimum; quad-core or better recommended for simultaneous cloning and parsing.'),
('Memory','4 GB minimum for local use; 8 GB or more recommended for development and larger repository scans.'),
('Storage','At least 2 GB free for application, dependencies, temporary shallow clones, logs and build output; more for concurrent workloads.'),
('Network / graphics','Reliable internet access for GitHub and OAuth services; no dedicated GPU is required.')], [1.7,4.8],8.5)
heading(doc,'2.6.3 Constraints',3)
bullet(doc,'Analysis is limited to validated public GitHub HTTPS repository URLs in the current public endpoint.')
bullet(doc,'Repository cloning requires Git and can be affected by network availability, upstream GitHub access, repository size, and server disk capacity.')
bullet(doc,'Configurable limits protect responsiveness: clone timeout, shallow depth, maximum repository size, parsed file count, dependency-source file count, and individual source-file size.')
bullet(doc,'The system depends on MongoDB availability; production startup is constrained by required database indexes and secure environment configuration.')
bullet(doc,'Security constraints include JWT secrets, encrypted provider tokens, secure cookie/session handling, CORS restrictions, rate limits, brute-force lockouts, and owned-resource authorization.')
bullet(doc,'AI outputs must be evidence-grounded and treated as recommendations requiring developer review; an LLM must not autonomously modify a repository.')
heading(doc,'2.6.4 Functional Requirements',3)
table(doc,['Requirement ID','Requirement description'],[
('CP_FR1','The system shall allow users to create accounts, verify email, sign in, refresh sessions, sign out, update profiles/settings, and reset passwords.'),
('CP_FR2','The system shall protect repository operations so a user can access only repositories and evidence records owned by that user.'),
('CP_FR3','The system shall accept and validate a supported public GitHub repository URL before analysis.'),
('CP_FR4','The system shall inspect, shallow-clone, parse and clean up a repository while enforcing configured safety limits.'),
('CP_FR5','The system shall extract and store repository metadata, file records, commit history, documentation records and basic dependency edges.'),
('CP_FR6','The system shall provide paginated retrieval of repository files, commits, dependencies and documentation, plus contributor aggregation and manifest parsing.'),
('CP_FR7','The system shall present repository intelligence and analytical dashboard views, including pipeline state and report-ready evidence.'),
('CP_FR8','The system shall support connected GitHub/GitLab source accounts without exposing provider access tokens to the frontend.'),
('CP_FR9','The planned analytics layer shall compute and expose debt, drift, risk and AI recommendations through protected backend interfaces.')], [1.05,5.45],8.4)
caption(doc,'Table 2.1 Functional Requirements')
heading(doc,'2.6.5 Non-Functional Requirements',3)
table(doc,['Requirement ID','Requirement description'],[
('CP_NFR1 Performance','The system shall use shallow clones, limits on repository size/files/dependency sources, pagination, and temporary-workspace cleanup to keep interactive analysis bounded.'),
('CP_NFR2 Security','The system shall use protected API routes, secure sessions/cookies, JWT access tokens, bcrypt password hashes, encrypted OAuth tokens, CORS controls, security headers, rate limits and brute-force protection.'),
('CP_NFR3 Reliability','The system shall represent analysis state, persist structured records, clean temporary clones best-effort, and ensure a failure in a bounded scan does not expose partial or unauthorized data as complete results.'),
('CP_NFR4 Usability','The product shall provide clear responsive screens, accessible controls, readable severity labels, empty states, and theme/density preferences.'),
('CP_NFR5 Maintainability','The system shall keep frontend, API, feature, service and database responsibilities separated, with documented route contracts and testable service modules.'),
('CP_NFR6 Scalability','The architecture shall keep analysis and planned analytical services behind API boundaries so execution can be expanded from synchronous local scans to scalable worker-based processing.'),
('CP_NFR7 Privacy','The system shall store repository metadata rather than retained clones; provider tokens shall remain encrypted and backend-only.')], [1.15,5.35],8.4)
caption(doc,'Table 2.2 Non-Functional Requirements')
heading(doc,'2.7 Block Diagram and Conceptual Model',2)
diagram(doc,'Figure 2.1 CodePulse System Block Diagram',[
'[Authenticated User / Browser]', '            ↓', '[React + Vite Presentation Layer]', '            ↓ HTTPS / API', '[Express API: Auth, Repository Intelligence, Integrations]', '            ↓', '[Repository Analysis Services]  ←  [GitHub / GitLab / Git]', '            ↓', '[MongoDB: Accounts, Repositories, Evidence Records]', '            ↓', '[Planned: Debt + Drift + Risk + AI Explainability]', '            ↓', '[Dashboard and Print-ready Reports]'])
para(doc,'The allowable sequence is: authenticate or establish a session; submit/select an owned repository; validate and analyze it; persist and retrieve evidence; present the evidence in dashboard or report views; and, when analytical services are enabled, compute and explain prioritized findings. A user cannot bypass ownership checks to access another user’s data.')

page_break(doc)
# CHAPTER 3
heading(doc,'CHAPTER 3: SYSTEM DESIGN',1)
para(doc,'This chapter presents the CodePulse architecture, modules, data design, interfaces, workflows, and reports.')
heading(doc,'3.1 System Architecture',2)
para(doc,'CodePulse follows a web-based three-tier architecture. The presentation tier is a React/Vite single-page application. The application tier is an Express API with security middleware and feature/service modules. The data tier uses MongoDB to store accounts, sessions, repository records, and analysis evidence. External integrations include GitHub/GitLab OAuth, GitHub repository metadata/raw manifests, and the Git executable for controlled repository cloning. The planned AI Explainability Engine consumes structured analytical context after debt, drift and risk stages produce evidence.')
diagram(doc,'Figure 3.1 Three-tier Architecture of CodePulse',[
'PRESENTATION TIER', '[React Dashboard | Auth | Account | Reports]', '                ↓', 'APPLICATION TIER', '[Express API | Auth | Repository Intelligence | Integrations]', '[Planned: Drift | Debt | Risk | AI Explainability]', '                ↓', 'DATA / INTEGRATION TIER', '[MongoDB Collections]  [GitHub/GitLab OAuth]  [Git / GitHub Raw Content]'])
heading(doc,'3.2 Module Design',2)
table(doc,['Module','Responsibilities'],[
('Authentication and Account Management','Signup, email verification, sign-in/sign-out, access/refresh token lifecycle, password reset, profile/settings, OAuth identity linking.'),
('Repository Intelligence','Validate GitHub URLs; inspect metadata; shallow clone; enumerate files; extract documentation, commits and dependencies; persist and remove temporary clone.'),
('Repository Query Services','List owned repositories and retrieve detail, files, commits, dependencies, documentation, contributors and manifest dependencies with pagination.'),
('Integration Services','Connect GitHub/GitLab accounts and list accessible repository sources while retaining provider tokens server-side.'),
('Dashboard and Reports','Orchestrate frontend data retrieval and display repository overview, pipeline, evidence panels, risk/debt/drift-oriented views and persisted report presentation.'),
('Analytical Engines (planned)','Calculate documentation drift, technical debt, knowledge debt and repository risk using stored evidence.'),
('AI Explainability (planned)','Construct targeted context from metrics, AST-like structure, documentation and history; generate evidence-backed explanations and action plans.')], [1.75,4.75],8.7)
heading(doc,'3.3 Data Flow Diagram',2)
heading(doc,'Level 0: Context Diagram',3)
diagram(doc,'Figure 3.2 DFD Level 0',[
'[User] → submit repository / view results → [CodePulse]', '[CodePulse] → dashboard, reports, account status → [User]', '[CodePulse] ↔ repository metadata / OAuth / raw manifests ↔ [GitHub / GitLab]', '[CodePulse] ↔ structured evidence ↔ [MongoDB]'])
heading(doc,'Level 1: Repository Analysis',3)
diagram(doc,'Figure 3.3 DFD Level 1',[
'[User] → (1. Authenticate) → [Auth + Session Store]', '[User] → (2. Submit URL) → [URL Validation + Metadata Check]', '                                 ↓', '                         [Clone + Parse Services]', '                                 ↓', '     [Files] [Docs] [Commits] [Dependencies] → [Repository Evidence Store]', '                                 ↓', '                       [Dashboard Query / Report View] → [User]'])
heading(doc,'Level 2: Planned Analytical Flow',3)
diagram(doc,'Figure 3.4 DFD Level 2',[
'[Stored Repository Evidence]', '        ↓', '(3.1 Debt Metrics) + (3.2 Documentation Drift) + (3.3 Knowledge Debt)', '        ↓', '(3.4 Risk Aggregation) → [Prioritized Findings]', '        ↓', '(3.5 Context Assembly + LLM) → [Explanation / Remediation Plan]', '        ↓', '[Dashboard + Report] → [User Review]'])
heading(doc,'3.4 ER Diagram',2)
diagram(doc,'Figure 3.5 Entity Relationship Diagram',[
'[users] 1 ──< [auth_sessions]', '[users] 1 ──< [email_verification_tokens] / [password_reset_tokens]', '[users] 1 ──< [oauth_accounts]', '[users] 1 ──< [repositories]', '[repositories] 1 ──< [repo_files]', '[repositories] 1 ──< [commits]', '[repositories] 1 ──< [dependencies]', '[repositories] 1 ──< [documentation]', '[repositories] 1 ──< [drift_findings]'])
heading(doc,'3.5 Database Design',2)
para(doc,'MongoDB is used as the operational data store. The design separates user/account data from repository evidence. A repository record owns its files, commits, dependencies, documentation and future drift findings; queries enforce the parent repository’s user ownership. Re-scanning the same user/repository updates the repository record and replaces its related evidence records, reducing stale duplicate data.')
heading(doc,'3.5.1 Table Design (MongoDB Collections)',3)
table(doc,['Collection','Key fields / datatype','Description and constraints'],[
('users','id: string; email: string; password_hash: string|null; email_verified: boolean','Account master. Email is normalized and unique; password hash may be null for OAuth-only accounts.'),
('auth_sessions','user_id: reference; token_hash: string; expires_at: date','Refresh-token sessions. Token is stored as a hash; revoked_at supports invalidation.'),
('oauth_accounts','provider: string; provider_user_id: string; user_id: reference; encrypted token: string','External identity link. Provider access token is AES-256-GCM encrypted and backend-only.'),
('repositories','user_id: reference; repo_url: string; status: enum; totals: integer','Owned repository metadata and lifecycle status: queued, running, completed or failed.'),
('repo_files','repository_id: reference; file_path: string; language: string; size: integer','Parsed repository file metadata.'),
('commits','repository_id: reference; commit_hash: string; author: string; commit_date: date','Commit history used for activity, churn and contributor analysis.'),
('dependencies','repository_id: reference; source_file: string; target_file: string; resolved: boolean','Basic file-to-file import/require dependency edge.'),
('documentation','repository_id: reference; doc_path: string; type: string; content: string','Extracted documentation text and capped content summary.'),
('drift_findings','repository_id: reference; drift_type: string; severity: enum; evidence: object','Future analytical findings for missing, outdated, incorrect or dead documentation.')], [1.35,2.6,2.5],7.8)
heading(doc,'3.5.2 Data Integrity and Constraints',3)
bullet(doc,'User email is lowercased and unique; credential fields are validated before account operations.')
bullet(doc,'Sensitive secret material is not stored in plain text: passwords, refresh tokens, verification tokens and reset tokens are hashed; OAuth access tokens are encrypted.')
bullet(doc,'Protected routes require a valid access token and repository queries constrain results to the authenticated owner.')
bullet(doc,'Repository URL validation restricts the current public analyzer to supported GitHub HTTPS URLs.')
bullet(doc,'Repository-analysis guards reject or stop work beyond configured capacity limits; temporary clone content is removed after analysis.')
bullet(doc,'Reference fields associate evidence records with a parent repository; a repository delete cascades to associated files, commits, dependencies and documentation.')
heading(doc,'3.5.3 Data Dictionary',3)
table(doc,['Table / field','Type / range','Default / validation','Meaning'],[
('users.email','string, email format','lowercase; unique','Account identifier.'),
('users.email_verified','boolean','false','Verification state.'),
('repositories.status','enum','queued | running | completed | failed','Analysis lifecycle state.'),
('repositories.total_files','integer, ≥ 0','0','Count of parsed files.'),
('repo_files.size','integer, ≥ 0','0','Source file size in bytes.'),
('commits.commit_hash','string','non-empty Git hash','Commit identity.'),
('dependencies.resolved','boolean','false','Whether an import maps to an internal file.'),
('documentation.truncated','boolean','false','Whether stored content was capped.'),
('drift_findings.severity','enum','Low | Medium | High | Critical','Finding priority.'),
('auth_sessions.expires_at','date/time','required','Session expiry threshold.')], [1.85,1.55,1.55,1.5],7.8)
heading(doc,'3.6 Interface and Procedural Design',2)
heading(doc,'3.6.1 User Interface Design',3)
para(doc,'The primary user is a developer or engineering manager working in a web browser, usually while evaluating an unfamiliar repository or planning maintenance work. The interface maps this task to a progression from sign-in, to repository selection/submission, to analysis progress, to evidence exploration and printable reporting. Responsive layouts, semantic colour tokens, keyboard-accessible controls, text labels for severity, theme selection and density settings support usability across devices and visual preferences.')
for title, lines in [
('Figure 3.6 Login and Sign-up Screen',['[CodePulse mark]','Email __________________________','Password _______________________','[ Sign in ]   [ Continue with GitHub ]','Forgot password?  Create account']),
('Figure 3.7 Repository Submission Screen',['Repository Intelligence','GitHub repository URL','https://github.com/owner/repository','[ Analyze repository ]','Recent repositories / source picker']),
('Figure 3.8 Dashboard Overview Screen',['Top bar: Repository selector | Theme | Account','Health / Technical Debt / Knowledge Debt / Risk KPI strip','Pipeline status     Repository intelligence','Trend chart         Recent evidence / recommendations']),
('Figure 3.9 Repository Intelligence Screen',['Repository metadata | Files | Manifests','Documentation list | Commit activity','Dependency explorer / contributor summary','Paginated evidence controls']),
('Figure 3.10 Risk and AI Workspace',['Risk heatmap / ranked modules','Drift findings queue','Evidence-backed recommendation cards','Status: planned analytics show empty/state messaging until available']),
('Figure 3.11 Account and Settings Screen',['Profile: name, company, title, timezone','Preferences: theme, density, scan frequency','Connected sources: GitHub / GitLab','Security: session / password controls'])]:
    diagram(doc,title,lines)
heading(doc,'Procedural Design and Business Rules',3)
table(doc,['Process','Rule / pseudocode'],[
('Analyze repository','IF user is authenticated AND URL is a valid public GitHub HTTPS URL, THEN check size/limits → shallow clone → parse → persist → cleanup; ELSE return a meaningful error.'),
('Read repository evidence','IF repository exists AND repository.user_id equals authenticated user id, THEN return requested paginated records; ELSE return not found/forbidden-safe response.'),
('Handle secrets','Never return password, refresh token, verification token, reset token or provider access token in API payloads. Store only hashes/encrypted values as applicable.'),
('Create recommendations (planned)','Gather only relevant metrics, documentation, history and code structure; request structured output; display recommendation as advisory evidence for human review.')], [1.75,4.75],8.7)
heading(doc,'3.7 Reports Design',2)
para(doc,'CodePulse is designed to present report-ready repository evidence. Reports are intended for engineering reviews, maintenance planning and stakeholder communication. Each report uses the selected repository as the primary input and filters its output by time, module, severity or evidence type where appropriate.')
table(doc,['Report','Inputs','Output fields'],[
('Repository Health Summary','Repository; scan date; optional period','Repository name, scan status, file/commit/documentation/dependency totals, health/debt/risk summaries when analytics are available.'),
('Documentation and Drift Report','Repository; severity; documentation area; period','Documentation paths, coverage categories, drift type, severity, evidence, age, suggested update (planned analytics output).'),
('Technical Risk and Remediation Report','Repository; module; risk threshold; time period','Module path, complexity, duplication, churn, dependency signals, risk level, explanation, priority and remediation steps (planned analytics output).')], [1.85,2.0,2.65],8.2)
for title, lines in [
('Figure 3.12 Report Screen: Repository Health Summary',['Repository Health Summary','Repository: owner/repository     Scan: latest','KPI tiles: files | commits | docs | dependencies','Health trend / pipeline status / evidence highlights','[ Print / Export ]']),
('Figure 3.13 Report Screen: Documentation and Drift',['Documentation and Drift Report','Filters: severity | documentation area | time period','Finding table: path | type | severity | evidence','Coverage chart / recommended documentation actions']),
('Figure 3.14 Report Screen: Technical Risk and Remediation',['Technical Risk and Remediation','Filters: module | threshold | period','Ranked risk table / dependency context','Recommendation: explanation | impact | effort | steps'])]:
    diagram(doc,title,lines)
para(doc,'The report layout is intentionally evidence-led: inputs determine the repository and filters; output fields make the underlying findings traceable; and AI-derived text, when enabled, is separated as an advisory explanation rather than treated as an unverified system fact.')

doc.core_properties.title='CodePulse Project Report - Chapters 1 to 3'
doc.core_properties.subject='Introduction, System Analysis and Requirements, and System Design'
doc.core_properties.author='CodePulse Project Team'
doc.save(OUT)
print(OUT)
